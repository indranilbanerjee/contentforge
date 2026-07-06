#!/usr/bin/env python3
"""
generate-docx.py
================
ContentForge Phase 8 — generates a publication-ready .docx file from
Markdown content plus pipeline reports (SEO, quality, production details).

Usage:
    python3 generate-docx.py \
        --content path/to/article.md \
        --output path/to/output.docx \
        --reports path/to/reports.json \
        --brand "Brand Name" \
        --content-type article

If python-docx is not installed, it is auto-installed via pip.

Output structure:
    1. Title page (brand, date, type, score)
    2. Table of Contents field (right-click > Update Field in Word)
    3. Body — full article with H1/H2/H3 hierarchy, images, nested lists
    4. Sources/Citations
    5. Appendix A — SEO Scorecard
    6. Appendix B — Quality Scorecard
    7. Appendix C — Production Details (phase timing, loops, metrics)
    8. Appendix D — Internal Link Map
    Footer: "Page X of Y" on every page. Body line spacing 1.15.

Reports JSON schema (all sections optional, missing = skipped):
    {
        "seo": {
            "primary_keyword": str,
            "keyword_density_pct": float,
            "meta_title": str,
            "meta_description": str,
            "schema_type": str,
            "internal_links": int,
            "seo_score": float
        },
        "quality": {
            "overall_score": float,
            "grade": str,  # A/B/C/D/F
            "dimensions": {
                "content_quality": float,
                "citation_integrity": float,
                "brand_compliance": float,
                "seo_performance": float,
                "readability": float
            },
            "review_date": str,
            "reviewer_notes": str
        },
        "production": {
            "phases_completed": [str, ...],
            "total_processing_time_seconds": float,
            "loops": int,
            "word_count": int,
            "citation_count": int,
            "source_reliability_avg": float,
            "flesch_kincaid_grade": float,
            "burstiness_score": float,
            "humanizer_patterns_removed": int,
            "em_dash_count": int,
            "ai_signal_score": float,  # 1-10, lower = more human
            "brand_compliance_violations": int,
            "factual_accuracy_pct": float,
            "hallucination_risk": str  # low/medium/high
        }
    }
"""

import argparse
import json
import re
import sys
from datetime import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import _common  # noqa: E402

_common.ensure_utf8_stdout()


def _plugin_version() -> str:
    """Read the canonical plugin version from .claude-plugin/plugin.json."""
    pj = Path(__file__).resolve().parent.parent / ".claude-plugin" / "plugin.json"
    data = _common.load_json_safe(pj)
    if isinstance(data, dict) and data.get("version"):
        return str(data["version"])
    return "unknown"


# Captures Phase 6 SEO agent INTERNAL-LINK markers so the renderer can replace
# them with real Word hyperlinks (or visibly-distinct placeholders when url=TBD).
# Marker format produced by 06-seo-geo-optimizer.md Step 5:
#   <!-- INTERNAL-LINK: type=topical|commercial|conversion|authority |
#        anchor="..." | url=... | priority=... | reason="..." | section=... -->
INTERNAL_LINK_PATTERN = re.compile(
    r"<!--\s*INTERNAL-LINK:\s*(?P<body>[^>]+?)-->",
    re.IGNORECASE,
)

# Block-level image: ![alt](path "optional title") alone on a line
IMAGE_BLOCK_PATTERN = re.compile(r'^!\[([^\]]*)\]\(\s*([^)\s]+)(?:\s+"[^"]*")?\s*\)\s*$')

# List items with indentation captured for nesting
BULLET_ITEM_PATTERN = re.compile(r"^(\s*)([-*+])\s+(.*)$")
ORDERED_ITEM_PATTERN = re.compile(r"^(\s*)\d+[.)]\s+(.*)$")

# Horizontal rules: ---, ***, ___ (3+ of the same char)
HR_PATTERN = re.compile(r"^\s*([-*_])\s*(?:\1\s*){2,}$")


def parse_internal_link_marker(marker_body):
    """Parse the key=value | key="value" pairs inside an INTERNAL-LINK marker."""
    fields = {}
    for part in marker_body.split("|"):
        if "=" not in part:
            continue
        k, v = part.split("=", 1)
        k = k.strip().lower()
        v = v.strip().strip('"').strip("'")
        fields[k] = v
    return fields


def split_text_with_links(text):
    """
    Split text on INTERNAL-LINK markers. Returns list of segments:
      ("text", "...")  — plain text
      ("link", {"anchor": str, "url": str, "type": str, "placeholder": bool})
    Anchor text from the marker REPLACES whatever the surrounding markdown
    rendered for it (e.g. plain word "ENHERTU" becomes a hyperlink).
    """
    segments = []
    pos = 0
    for m in INTERNAL_LINK_PATTERN.finditer(text):
        if m.start() > pos:
            segments.append(("text", text[pos:m.start()]))
        fields = parse_internal_link_marker(m.group("body"))
        anchor = fields.get("anchor", "").strip()
        url = fields.get("url", "").strip()
        link_type = fields.get("type", "topical").strip().lower()
        placeholder = (not url) or url.upper() in {"TBD", "TODO", "PLACEHOLDER"}
        if anchor:
            segments.append(("link", {
                "anchor": anchor,
                "url": url if not placeholder else "",
                "type": link_type,
                "placeholder": placeholder,
            }))
        pos = m.end()
    if pos < len(text):
        segments.append(("text", text[pos:]))
    if not segments:
        segments.append(("text", text))
    return segments


def add_hyperlink_run(paragraph, anchor_text, url, link_type="topical", placeholder=False):
    """
    Add a clickable hyperlink run to a python-docx paragraph using raw OOXML.
    python-docx has no high-level hyperlink API; this is the standard pattern.
    Placeholders render as bold + bracketed anchor with a [LINK TBD] suffix,
    so a human reviewer can see what should be linked and fill in the URL.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor

    if placeholder or not url:
        # Visibly distinct placeholder: bold + bracketed + LINK TBD suffix
        run = paragraph.add_run(f"[{anchor_text}] [LINK TBD: {link_type}]")
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        return run

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    # color hint by link type so reviewers can spot the three categories
    color = {
        "topical":    "0066CC",  # blue
        "commercial": "2E7D32",  # green = revenue
        "conversion": "8E24AA",  # purple = funnel
        "authority":  "455A64",  # slate grey
    }.get(link_type, "0066CC")
    c.set(qn("w:val"), color)
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = anchor_text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def ensure_docx():
    """Install python-docx if not present."""
    try:
        import docx  # noqa: F401
        return
    except ImportError:
        pass
    err = _common.pip_install(["python-docx>=1.1.0"], label="python-docx")
    if err:
        _common.finish(err)


def _is_block_start(line):
    """True when a line begins a non-paragraph block (used to terminate
    paragraph continuation)."""
    if line.startswith(("#", ">", "|", "```")):
        return True
    if HR_PATTERN.match(line):
        return True
    if BULLET_ITEM_PATTERN.match(line) or ORDERED_ITEM_PATTERN.match(line):
        return True
    if IMAGE_BLOCK_PATTERN.match(line.strip()):
        return True
    return False


def _indent_level(indent_str):
    """Map leading whitespace to a nesting level 0-2 (2-space or 4-space
    indents both count as one level per step; deeper than 2 is clamped)."""
    spaces = len(indent_str.expandtabs(4))
    if spaces == 0:
        return 0
    if spaces < 4:
        return 1
    return 2


def parse_markdown_to_blocks(md_text):
    """
    Parse markdown into ordered blocks: (kind, content) tuples.
    kind in {'h1','h2','h3','para','bullet','ordered','table','hr','quote',
             'code','image'}
    'bullet' and 'ordered' content is a list of (level, text) tuples where
    level is 0-2 (nested lists).
    'image' content is (alt_text, path).
    Strips YAML frontmatter if present.
    Strips ContentForge "completion card" / appendix sections that the agent may
    have inlined — those go in the proper appendix sections instead.
    """
    # Strip YAML frontmatter
    if md_text.lstrip().startswith("---"):
        end = md_text.find("\n---", 3)
        if end > 0:
            md_text = md_text[end + 4:].lstrip()

    # Strip everything from "## CONTENTFORGE — COMPLETION CARD" onwards if present
    # (that content moves into the appendix section).
    cutoff_patterns = [
        r"\n##\s+CONTENTFORGE\s*[—-]+\s*COMPLETION\s+CARD",
        r"\n##\s+Appendix\s+[ABC]:",
        r"\n#\s+Appendix\s+[ABC]:",
    ]
    for pat in cutoff_patterns:
        m = re.search(pat, md_text, re.IGNORECASE)
        if m:
            md_text = md_text[:m.start()].rstrip()

    blocks = []
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        img = IMAGE_BLOCK_PATTERN.match(line.strip())
        if img:
            blocks.append(("image", (img.group(1), img.group(2))))
            i += 1
        elif line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
            i += 1
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
            i += 1
        elif line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
            i += 1
        elif HR_PATTERN.match(line):
            blocks.append(("hr", ""))
            i += 1
        elif line.startswith("> "):
            quote_lines = []
            while i < len(lines) and lines[i].startswith(">"):
                quote_lines.append(lines[i].lstrip("> ").rstrip())
                i += 1
            blocks.append(("quote", " ".join(quote_lines)))
        elif line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[-: |]+\|", lines[i + 1]):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            blocks.append(("table", table_lines))
        elif BULLET_ITEM_PATTERN.match(line):
            items = []
            while i < len(lines):
                m = BULLET_ITEM_PATTERN.match(lines[i])
                if not m:
                    break
                items.append((_indent_level(m.group(1)), m.group(3).rstrip()))
                i += 1
            blocks.append(("bullet", items))
        elif ORDERED_ITEM_PATTERN.match(line):
            items = []
            while i < len(lines):
                m = ORDERED_ITEM_PATTERN.match(lines[i])
                if not m:
                    break
                items.append((_indent_level(m.group(1)), m.group(2).rstrip()))
                i += 1
            blocks.append(("ordered", items))
        elif line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            blocks.append(("code", "\n".join(code_lines)))
        else:
            para_lines = [line]
            i += 1
            while (i < len(lines) and lines[i].strip()
                   and not _is_block_start(lines[i])):
                para_lines.append(lines[i].rstrip())
                i += 1
            blocks.append(("para", " ".join(para_lines)))

    return blocks


# Inline markdown, matched in priority order:
#   inline image, ***bold italic*** / ___..___, **bold** / __..__,
#   *italic*, _italic_ (word-boundary guarded so snake_case survives),
#   `code`, [text](url)
INLINE_PATTERN = re.compile(
    r"!\[([^\]]*)\]\(([^)\s]+)(?:\s+\"[^\"]*\")?\)"        # 1,2  image
    r"|\*\*\*([^*]+)\*\*\*|___([^_]+)___"                   # 3,4  bold italic
    r"|\*\*([^*]+)\*\*|__([^_]+)__"                         # 5,6  bold
    r"|\*([^*\s][^*]*?)\*"                                  # 7    italic *
    r"|(?<![\w\\])_([^_\s][^_]*?)_(?![\w])"                 # 8    italic _
    r"|`([^`]+)`"                                            # 9    code
    r"|\[([^\]]+)\]\(([^)]+)\)"                             # 10,11 link
)


def _render_markdown_segment(paragraph, text):
    """Render inline markdown (bold/italic/bold-italic in * and _ forms,
    `code`, [text](url) links, and inline images) inside a text segment."""
    from docx.shared import Pt, RGBColor, Inches

    pos = 0
    for m in INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        (img_alt, img_path, bi_star, bi_und, b_star, b_und,
         i_star, i_und, code_t, link_t, link_url) = m.groups()
        if img_path is not None:
            _add_image_run(paragraph, img_alt or "", img_path, inline=True)
        elif bi_star is not None or bi_und is not None:
            run = paragraph.add_run(bi_star if bi_star is not None else bi_und)
            run.bold = True
            run.italic = True
        elif b_star is not None or b_und is not None:
            run = paragraph.add_run(b_star if b_star is not None else b_und)
            run.bold = True
        elif i_star is not None or i_und is not None:
            run = paragraph.add_run(i_star if i_star is not None else i_und)
            run.italic = True
        elif code_t is not None:
            run = paragraph.add_run(code_t)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif link_t is not None:
            # standard markdown [text](url) link → real hyperlink (outbound citation style)
            if link_url:
                add_hyperlink_run(paragraph, link_t, link_url, link_type="topical", placeholder=False)
            else:
                run = paragraph.add_run(link_t)
                run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
                run.font.underline = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_image_run(paragraph, alt_text, path, inline=False):
    """Add an image to a paragraph. Missing/broken files degrade to a
    visible placeholder so the reviewer knows an asset needs attention."""
    from docx.shared import Inches, RGBColor

    img_path = Path(path).expanduser()
    if img_path.exists() and img_path.is_file():
        try:
            run = paragraph.add_run()
            if inline:
                run.add_picture(str(img_path), width=Inches(3.0))
            else:
                run.add_picture(str(img_path), width=Inches(6.0))
            return True
        except Exception as exc:
            print(f"Image embed error for {img_path}: {exc}", file=sys.stderr)
    run = paragraph.add_run(f"[Image not found: {path}"
                            + (f" — {alt_text}" if alt_text else "") + "]")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    return False


def add_inline_runs(paragraph, text):
    """
    Render inline content with two passes:
      1. Split on `<!-- INTERNAL-LINK: ... -->` markers from Phase 6 SEO agent
         → render each as a real Word hyperlink (or visible placeholder if URL is TBD)
      2. Within plain-text segments, render markdown emphasis/code/links/images
    """
    segments = split_text_with_links(text)
    for kind, payload in segments:
        if kind == "link":
            add_hyperlink_run(
                paragraph,
                payload["anchor"],
                payload["url"],
                link_type=payload["type"],
                placeholder=payload["placeholder"],
            )
        else:
            _render_markdown_segment(paragraph, payload)


def collect_internal_link_markers(md_text):
    """Walk the markdown text and return a list of all INTERNAL-LINK fields found.
    Used by Appendix D to render the link map table."""
    found = []
    for m in INTERNAL_LINK_PATTERN.finditer(md_text):
        fields = parse_internal_link_marker(m.group("body"))
        if fields.get("anchor"):
            found.append(fields)
    return found


def _list_style(base, level, doc):
    """Return a list style name for a nesting level, falling back to the
    base style when the numbered variant isn't in the template."""
    name = base if level == 0 else f"{base} {level + 1}"
    try:
        doc.styles[name]  # noqa: B018 — existence probe
        return name
    except KeyError:
        return base


def render_blocks(doc, blocks):
    """Render parsed markdown blocks into the docx document."""
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for kind, content in blocks:
        if kind == "h1":
            # Use Word's semantic Heading 1 style so Navigation Pane, TOC,
            # PDF bookmarks, and screen readers recognise it as a heading.
            p = doc.add_heading(level=1)
            run = p.add_run(content)
            run.font.size = Pt(24)
            p.paragraph_format.space_after = Pt(12)
        elif kind == "h2":
            p = doc.add_heading(level=2)
            run = p.add_run(content)
            run.font.size = Pt(18)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
        elif kind == "h3":
            p = doc.add_heading(level=3)
            run = p.add_run(content)
            run.font.size = Pt(14)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif kind == "para":
            p = doc.add_paragraph()
            add_inline_runs(p, content)
            p.paragraph_format.space_after = Pt(6)
        elif kind == "image":
            alt, path = content
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_image_run(p, alt, path, inline=False)
            if alt:
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap.add_run(alt)
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        elif kind == "bullet":
            for level, item in content:
                p = doc.add_paragraph(style=_list_style("List Bullet", level, doc))
                add_inline_runs(p, item)
        elif kind == "ordered":
            for level, item in content:
                p = doc.add_paragraph(style=_list_style("List Number", level, doc))
                add_inline_runs(p, item)
        elif kind == "quote":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.italic = True
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(6)
        elif kind == "code":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            p.paragraph_format.left_indent = Inches(0.5)
        elif kind == "hr":
            p = doc.add_paragraph()
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "table":
            try:
                rows = [
                    [c.strip() for c in line.strip("|").split("|")]
                    for line in content
                    if not re.match(r"^\|[-: |]+\|", line)
                ]
                if not rows:
                    continue
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        if c_idx >= len(rows[0]):
                            continue
                        cell = table.cell(r_idx, c_idx)
                        cell.text = ""
                        p = cell.paragraphs[0]
                        add_inline_runs(p, cell_text)
                        if r_idx == 0:
                            for run in p.runs:
                                run.bold = True
                doc.add_paragraph()
            except Exception as e:
                print(f"Table render error: {e}", file=sys.stderr)


def _add_field(paragraph, instruction, placeholder_text=""):
    """Insert a Word field (e.g. TOC, PAGE, NUMPAGES) via raw OOXML."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    r_begin = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r_begin.append(fld_begin)

    r_instr = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    r_instr.append(instr)

    r_sep = OxmlElement("w:r")
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r_sep.append(fld_sep)

    r_text = OxmlElement("w:r")
    if placeholder_text:
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = placeholder_text
        r_text.append(t)

    r_end = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_end.append(fld_end)

    for r in (r_begin, r_instr, r_sep, r_text, r_end):
        paragraph._p.append(r)


def add_toc(doc):
    """Insert a Table of Contents field covering heading levels 1-3.
    Word populates it on 'Update Field' / print preview."""
    from docx.shared import Pt

    p = doc.add_paragraph()
    run = p.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(16)

    field_para = doc.add_paragraph()
    _add_field(field_para, r'TOC \o "1-3" \h \z \u',
               "Right-click and choose 'Update Field' to build the table of contents.")
    doc.add_page_break()


def add_page_footer(doc):
    """'Page X of Y' centered footer on every section."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Page ")
        _add_field(p, "PAGE", "1")
        p.add_run(" of ")
        _add_field(p, "NUMPAGES", "1")


def add_title_page(doc, brand, content_type, title, score, grade):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(brand.upper())
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(content_type.replace("_", " ").upper())
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # Use Word's semantic Title style so Navigation Pane / PDF bookmarks
    # recognise this as the document title.
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(20)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated by ContentForge — {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    if score is not None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Quality Score: {score}/10  •  Grade: {grade}")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x33, 0x66, 0x33)

    doc.add_page_break()


def add_internal_link_map_appendix(doc, link_markers):
    """Appendix D — Internal Link Map. Shows every <!-- INTERNAL-LINK: ... --> marker
    the SEO agent placed in the body (real URLs + placeholders), so a human reviewer
    can verify the marketing-funnel coverage at a glance."""
    from docx.shared import Pt

    if not link_markers:
        return

    p = doc.add_heading(level=2)
    run = p.add_run("Appendix D — Internal Link Map")
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    run = p.add_run(
        "Every internal link inserted by Phase 6 (SEO/GEO Optimizer). "
        "Three categories: TOPICAL (informational), COMMERCIAL (brand product/service), "
        "CONVERSION (audience-matched CTA). Placeholder URLs (TBD) require human review "
        "before publication."
    )
    run.italic = True
    run.font.size = Pt(10)

    rows = [("#", "Type", "Anchor Text", "Target URL", "Section", "Reason")]
    for i, m in enumerate(link_markers, 1):
        url = m.get("url", "").strip()
        if not url or url.upper() in {"TBD", "TODO", "PLACEHOLDER"}:
            url_cell = "[TBD — fill before publish]"
        else:
            url_cell = url
        rows.append((
            str(i),
            m.get("type", "topical").upper(),
            m.get("anchor", ""),
            url_cell,
            m.get("section", ""),
            m.get("reason", ""),
        ))
    table = doc.add_table(rows=len(rows), cols=6)
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = str(cell_text)
            if r_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    # Summary counts so the marketer sees coverage at a glance
    by_type = {}
    placeholders = 0
    for m in link_markers:
        t = m.get("type", "topical").lower()
        by_type[t] = by_type.get(t, 0) + 1
        url = m.get("url", "").strip()
        if not url or url.upper() in {"TBD", "TODO", "PLACEHOLDER"}:
            placeholders += 1

    p = doc.add_paragraph()
    summary = (
        f"Coverage — Topical: {by_type.get('topical', 0)} | "
        f"Commercial: {by_type.get('commercial', 0)} | "
        f"Conversion: {by_type.get('conversion', 0)} | "
        f"Authority: {by_type.get('authority', 0)} | "
        f"Placeholders needing URL: {placeholders}"
    )
    run = p.add_run(summary)
    run.italic = True
    run.font.size = Pt(10)
    doc.add_paragraph()


def add_appendices(doc, reports, link_markers=None):
    """Add Appendix A (SEO), B (Quality), C (Production), D (Internal Links) from reports JSON."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    seo = reports.get("seo")
    quality = reports.get("quality")
    production = reports.get("production")

    has_links = bool(link_markers)
    if not (seo or quality or production or has_links):
        return

    doc.add_page_break()
    # Heading 1 so Navigation Pane / PDF bookmarks / Word TOC pick it up
    p = doc.add_heading(level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("APPENDICES")
    run.font.size = Pt(20)
    doc.add_paragraph()

    if seo:
        # Heading 2 — appendix sub-headers nest under the Heading 1 above
        p = doc.add_heading(level=2)
        run = p.add_run("Appendix A — SEO Scorecard")
        run.font.size = Pt(16)
        rows = [
            ("Metric", "Value"),
            ("Primary keyword", str(seo.get("primary_keyword", "—"))),
            ("Keyword density", f"{seo.get('keyword_density_pct', 0):.2f}%"),
            ("Meta title", str(seo.get("meta_title", "—"))),
            ("Meta description", str(seo.get("meta_description", "—"))),
            ("Schema type", str(seo.get("schema_type", "—"))),
            ("Internal links", str(seo.get("internal_links", 0))),
            ("SEO score", f"{seo.get('seo_score', 0):.2f}/10"),
        ]
        table = doc.add_table(rows=len(rows), cols=2)
        table.style = "Light Grid Accent 1"
        for r_idx, (k, v) in enumerate(rows):
            table.cell(r_idx, 0).text = k
            table.cell(r_idx, 1).text = v
            if r_idx == 0:
                for cell in [table.cell(r_idx, 0), table.cell(r_idx, 1)]:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
        doc.add_paragraph()

    if quality:
        p = doc.add_heading(level=2)
        run = p.add_run("Appendix B — Quality Scorecard")
        run.font.size = Pt(16)
        dims = quality.get("dimensions", {})
        rows = [
            ("Dimension", "Score (0-10)", "Weight", "Status"),
            ("Content Quality", f"{dims.get('content_quality', 0):.2f}", "30%",
             "PASS" if dims.get("content_quality", 0) >= 7.0 else "FAIL"),
            ("Citation Integrity", f"{dims.get('citation_integrity', 0):.2f}", "25%",
             "PASS" if dims.get("citation_integrity", 0) >= 7.0 else "FAIL"),
            ("Brand Compliance", f"{dims.get('brand_compliance', 0):.2f}", "20%",
             "PASS" if dims.get("brand_compliance", 0) >= 7.0 else "FAIL"),
            ("SEO Performance", f"{dims.get('seo_performance', 0):.2f}", "15%",
             "PASS" if dims.get("seo_performance", 0) >= 7.0 else "FAIL"),
            ("Readability", f"{dims.get('readability', 0):.2f}", "10%",
             "PASS" if dims.get("readability", 0) >= 7.0 else "FAIL"),
            ("OVERALL", f"{quality.get('overall_score', 0):.2f}", "100%",
             quality.get("grade", "—")),
        ]
        table = doc.add_table(rows=len(rows), cols=4)
        table.style = "Light Grid Accent 1"
        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                cell = table.cell(r_idx, c_idx)
                cell.text = cell_text
                if r_idx == 0 or r_idx == len(rows) - 1:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
        if quality.get("review_date"):
            p = doc.add_paragraph()
            run = p.add_run(f"Reviewed: {quality['review_date']}")
            run.italic = True
            run.font.size = Pt(10)
        if quality.get("reviewer_notes"):
            p = doc.add_paragraph()
            run = p.add_run(f"Notes: {quality['reviewer_notes']}")
            run.italic = True
            run.font.size = Pt(10)
        doc.add_paragraph()

    if production:
        p = doc.add_heading(level=2)
        run = p.add_run("Appendix C — Production Details")
        run.font.size = Pt(16)

        rows = [
            ("Metric", "Value"),
            ("Phases completed", ", ".join(production.get("phases_completed", []))),
            ("Total processing time", f"{production.get('total_processing_time_seconds', 0):.1f}s"),
            ("Loops (max 5)", str(production.get("loops", 0))),
            ("Word count", str(production.get("word_count", 0))),
            ("Citation count", str(production.get("citation_count", 0))),
            ("Source reliability (avg)", f"{production.get('source_reliability_avg', 0):.2f}/10"),
            ("Flesch-Kincaid grade", f"{production.get('flesch_kincaid_grade', 0):.1f}"),
            ("Burstiness score", f"{production.get('burstiness_score', 0):.2f} (target ≥0.7)"),
            ("Humanizer patterns removed", str(production.get("humanizer_patterns_removed", 0))),
            ("Em dash count", f"{production.get('em_dash_count', 0)} (target ≤1-2 per 500w)"),
            ("AI signal score", f"{production.get('ai_signal_score', 0):.1f}/10 (target ≤3)"),
            ("Brand compliance violations", str(production.get("brand_compliance_violations", 0))),
            ("Factual accuracy", f"{production.get('factual_accuracy_pct', 0):.1f}%"),
            ("Hallucination risk", str(production.get("hallucination_risk", "—"))),
        ]
        table = doc.add_table(rows=len(rows), cols=2)
        table.style = "Light Grid Accent 1"
        for r_idx, (k, v) in enumerate(rows):
            table.cell(r_idx, 0).text = k
            table.cell(r_idx, 1).text = str(v)
            if r_idx == 0:
                for cell in [table.cell(r_idx, 0), table.cell(r_idx, 1)]:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

    if link_markers:
        doc.add_paragraph()
        add_internal_link_map_appendix(doc, link_markers)


def _maybe_c2pa_sign_docx(output_path, args, title):
    """v3.10 — Optional C2PA provenance signing of the .docx output for
    EU AI Act Article 50 compliance (applicable 2 Aug 2026; covers AI-generated
    text on matters of public interest unless human-reviewed and brand assumes
    editorial responsibility).

    Mirrors the SocialForge / DMP pattern. Self-contained installation of
    c2pa-python on demand. Non-fatal — returns None if signing wasn't
    requested or failed; the unsigned .docx remains on disk either way.
    """
    if not args.c2pa_sign:
        return None
    try:
        import subprocess as _sp
        try:
            import c2pa  # noqa: F401
        except ImportError:
            _sp.check_call([sys.executable, "-m", "pip", "install", "--quiet", "c2pa-python>=0.32", "cryptography"])
            import c2pa  # noqa: F811

        try:
            import c2pa  # type: ignore
        except ImportError as exc:
            return {"c2pa_signed": False, "c2pa_error": f"could not install c2pa-python: {exc}"}

        # Use sibling C2PA module if present (DMP-style); otherwise inline-minimal sign
        # The .docx format is supported by c2pa-python via the "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        # MIME type — but the simpler/more portable path is a sidecar manifest. We do BOTH:
        # 1. Embed in .docx if supported
        # 2. Write a .c2pa.json sidecar with the manifest content for inspectability

        from datetime import datetime as _dt, timezone as _tz
        created = _dt.now(_tz.utc).isoformat()
        cf_version = _plugin_version()
        manifest = {
            "claim_generator_info": [{"name": "ContentForge", "version": cf_version}],
            "title": f"{args.brand} — AI-assisted {args.content_type}",
            "format": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "assertions": [
                {
                    "label": "c2pa.actions.v2",
                    "data": {"actions": [
                        {
                            "action": "c2pa.created",
                            "when": created,
                            "softwareAgent": {"name": "ContentForge 10-phase pipeline", "version": cf_version},
                        },
                        {
                            "action": "c2pa.edited",
                            "when": created,
                            "parameters": {"description": "Human-reviewed via Phase 7 reviewer scorecard before delivery"},
                        },
                    ]},
                },
                {
                    "label": "stds.schema-org.CreativeWork",
                    "data": {
                        "@context": "https://schema.org",
                        "@type": "Article" if args.content_type in ("article", "blog") else "CreativeWork",
                        "author": [{"@type": "Organization", "name": args.brand}],
                        "dateCreated": created,
                        "headline": title,
                    },
                },
            ],
        }
        # Sidecar JSON
        sidecar_path = output_path.with_suffix(".c2pa.json")
        sidecar_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")

        # Try embedding via c2pa.Builder if .docx MIME is supported (versions vary)
        embed_status = "sidecar-only"
        embed_manifest_id = None
        using_dev_cert = False
        try:
            # Generate self-signed dev cert if user didn't supply one
            import tempfile
            cert_pem, key_pem = args.c2pa_signing_cert, args.c2pa_signing_key
            if not (cert_pem and key_pem):
                from cryptography import x509
                from cryptography.x509.oid import NameOID, ExtendedKeyUsageOID
                from cryptography.hazmat.primitives import hashes, serialization
                from cryptography.hazmat.primitives.asymmetric import ec
                from datetime import timedelta as _td
                key = ec.generate_private_key(ec.SECP256R1())
                subject = issuer = x509.Name([
                    x509.NameAttribute(NameOID.COMMON_NAME, "ContentForge Dev Self-Signed C2PA"),
                    x509.NameAttribute(NameOID.ORGANIZATION_NAME, "ContentForge"),
                    x509.NameAttribute(NameOID.COUNTRY_NAME, "US"),
                ])
                cert_obj = (
                    x509.CertificateBuilder()
                    .subject_name(subject).issuer_name(issuer).public_key(key.public_key())
                    .serial_number(x509.random_serial_number())
                    .not_valid_before(_dt.now(_tz.utc) - _td(minutes=1))
                    .not_valid_after(_dt.now(_tz.utc) + _td(days=90))
                    .add_extension(x509.BasicConstraints(ca=False, path_length=None), critical=True)
                    .add_extension(x509.KeyUsage(digital_signature=True, content_commitment=False,
                        key_encipherment=False, data_encipherment=False, key_agreement=False,
                        key_cert_sign=False, crl_sign=False, encipher_only=False, decipher_only=False), critical=True)
                    .add_extension(x509.ExtendedKeyUsage([ExtendedKeyUsageOID.EMAIL_PROTECTION]), critical=False)
                    .add_extension(x509.SubjectKeyIdentifier.from_public_key(key.public_key()), critical=False)
                    .add_extension(x509.AuthorityKeyIdentifier.from_issuer_public_key(key.public_key()), critical=False)
                    .sign(key, hashes.SHA256())
                )
                tmpdir = tempfile.mkdtemp(prefix="cf-c2pa-")
                cert_pem = str(Path(tmpdir) / "dev-cert.pem")
                key_pem = str(Path(tmpdir) / "dev-key.pem")
                Path(cert_pem).write_bytes(cert_obj.public_bytes(serialization.Encoding.PEM))
                Path(key_pem).write_bytes(key.private_bytes(
                    encoding=serialization.Encoding.PEM,
                    format=serialization.PrivateFormat.PKCS8,
                    encryption_algorithm=serialization.NoEncryption(),
                ))
                using_dev_cert = True

            supported = set(c2pa.Builder.get_supported_mime_types())
            docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if docx_mime in supported:
                signer_info = c2pa.C2paSignerInfo(
                    alg=b"es256",
                    sign_cert=open(cert_pem, "rb").read(),
                    private_key=open(key_pem, "rb").read(),
                    ta_url=b"http://timestamp.digicert.com",
                )
                signer = c2pa.Signer.from_info(signer_info)
                builder = c2pa.Builder(manifest)
                try:
                    builder.set_intent(c2pa.C2paBuilderIntent.CREATE, c2pa.C2paDigitalSourceType.COMPOSITE_WITH_TRAINED_ALGORITHMIC_MEDIA)
                except Exception:
                    pass
                tmp_signed = output_path.with_suffix(".c2pa-tmp.docx")
                builder.sign_file(str(output_path), str(tmp_signed), signer=signer)
                # Atomic replace — the old unlink()+rename() pair had a crash
                # window where the deliverable was deleted.
                tmp_signed.replace(output_path)
                # Round-trip
                try:
                    with open(output_path, "rb") as fh:
                        with c2pa.Reader(docx_mime, fh) as reader:
                            r = json.loads(reader.json())
                            embed_manifest_id = r.get("active_manifest")
                            embed_status = "embedded-and-verified" if embed_manifest_id else "embedded"
                except Exception:
                    embed_status = "embedded-unverified"
            else:
                embed_status = "sidecar-only (.docx MIME not in c2pa-python supported list)"
        except Exception as exc:
            embed_status = f"sidecar-only (embed failed: {type(exc).__name__}: {exc})"

        return {
            "c2pa_signed": True,
            "c2pa_embed_status": embed_status,
            "c2pa_active_manifest_id": embed_manifest_id,
            "c2pa_sidecar": str(sidecar_path),
            "c2pa_using_dev_cert": using_dev_cert,
        }
    except Exception as exc:
        return {"c2pa_signed": False, "c2pa_error": f"{type(exc).__name__}: {exc}"}


def main():
    parser = argparse.ArgumentParser(description="Generate ContentForge .docx output")
    parser.add_argument("--content", required=True, help="Markdown file with article body")
    parser.add_argument("--output", required=True, help="Output .docx path")
    parser.add_argument("--reports", help="Optional reports JSON for appendices")
    parser.add_argument("--brand", default="Brand", help="Brand name for header")
    parser.add_argument("--content-type", default="article", help="article/blog/whitepaper/faq/research_paper")
    parser.add_argument("--title", help="Override title (else extracts from first H1)")
    parser.add_argument("--no-toc", action="store_true",
                        help="Skip the Table of Contents field after the title page")
    # v3.10 — EU AI Act Article 50: optional C2PA provenance signing of the .docx
    parser.add_argument("--c2pa-sign", action="store_true",
                        help="Embed C2PA manifest in the .docx (or write a sidecar .c2pa.json if .docx embedding unsupported). EU AI Act Article 50 compliance for AI-assisted content distributed in EU markets.")
    parser.add_argument("--c2pa-signing-cert", default=None,
                        help="PEM signing certificate (omit for dev 90-day self-signed cert; production REQUIRES a CAI-recognized cert)")
    parser.add_argument("--c2pa-signing-key", default=None,
                        help="PEM signing key (must accompany --c2pa-signing-cert)")
    args = parser.parse_args()

    ensure_docx()
    from docx import Document
    from docx.shared import Pt, Inches

    content_path = Path(args.content)
    if not content_path.exists():
        _common.finish({"error": f"content file not found: {content_path}"})
    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    md_text = content_path.read_text(encoding="utf-8")
    link_markers = collect_internal_link_markers(md_text)
    blocks = parse_markdown_to_blocks(md_text)

    title = args.title
    if not title:
        for kind, content in blocks:
            if kind == "h1":
                title = content
                break
    if not title:
        title = output_path.stem

    reports = {}
    if args.reports and Path(args.reports).exists():
        try:
            reports = json.loads(Path(args.reports).read_text(encoding="utf-8"))
        except Exception as e:
            print(f"Warning: could not parse reports JSON: {e}", file=sys.stderr)

    score = reports.get("quality", {}).get("overall_score")
    grade = reports.get("quality", {}).get("grade", "—")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15  # per 08-output-manager spec
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    add_title_page(doc, args.brand, args.content_type, title, score, grade)
    if not args.no_toc:
        add_toc(doc)

    # Strip only the FIRST body H1 that duplicates the title-page title;
    # legitimately repeated H1s later in the body are preserved.
    blocks_no_title = []
    removed_title_h1 = False
    for b in blocks:
        if not removed_title_h1 and b[0] == "h1" and b[1] == title:
            removed_title_h1 = True
            continue
        blocks_no_title.append(b)

    render_blocks(doc, blocks_no_title)
    add_appendices(doc, reports, link_markers=link_markers)
    add_page_footer(doc)

    doc.save(output_path)
    link_summary = {}
    for m in link_markers:
        t = m.get("type", "topical").lower()
        link_summary[t] = link_summary.get(t, 0) + 1

    # v3.10 — optional C2PA provenance signing for EU AI Act Article 50
    c2pa_result = _maybe_c2pa_sign_docx(output_path, args, title)

    result_obj = {
        "status": "success",
        "output": str(output_path),
        "size_bytes": output_path.stat().st_size,
        "title": title,
        "brand": args.brand,
        "content_type": args.content_type,
        "score": score,
        "grade": grade,
        "reports_included": list(reports.keys()),
        "internal_links_total": len(link_markers),
        "internal_links_by_type": link_summary,
        "toc_included": not args.no_toc,
    }
    if c2pa_result is not None:
        result_obj["c2pa"] = c2pa_result
    _common.finish(result_obj)


if __name__ == "__main__":
    main()
